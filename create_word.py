# -*- coding: utf-8 -*-
from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, re, html

app = Flask(__name__)

# ---------------- فارسی‌ساز ----------------
class PersianTextProcessor:
    def clean_text(self, text):
        if not text:
            return ''
        # جایگزینی حروف عربی با فارسی
        text = text.replace('ي', 'ی').replace('ك', 'ک').replace('ە', 'ه').replace('ؤ', 'و')
        text = text.replace('\x00', '')  # حذف NULL عجیب
        # یکنواخت‌سازی فاصله‌ها و علائم
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\s+([.,،؛:!؟»\\)])', r'\1', text)
        text = re.sub(r'([(«])\s+', r'\1', text)
        # حذف نمادهای markdown (ستاره و اندرلاین تودرتو)
        text = re.sub(r'(\*\*|__)(.*?)\1', r'\2', text)
        return text.strip()

# ---------------- سازنده سند ----------------
class SmartDocumentGenerator:
    def __init__(self):
        self.doc = Document()
        self.text_processor = PersianTextProcessor()
        self._setup_doc()

    def _setup_doc(self):
        s = self.doc.sections[0]
        s.page_height = Inches(11.69)
        s.page_width = Inches(8.27)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

    def _set_rtl(self, p):
        pPr = p._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def _set_cell_borders(self, cell):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def _set_cell_shading(self, cell, is_header=False):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3' if is_header else 'FFFFFF')
        tcPr.append(shading)

    def _set_cell_margins(self, cell):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '100')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)

    def _parse_inline_marks(self, text):
        """
        شناسایی بخش‌های bold/underline با فرمت **bold** یا __underline__
        ولی خود علامت‌ها حذف می‌شوند.
        """
        parts = []
        pattern = r'(\*\*(.*?)\*\*|__(.*?)__)'
        last_end = 0
        for m in re.finditer(pattern, text):
            if m.start() > last_end:
                parts.append({'text': text[last_end:m.start()], 'bold': False, 'underline': False})
            inner_text = m.group(2) if m.group(2) else m.group(3)
            if m.group(2):  # bold
                parts.append({'text': inner_text, 'bold': True, 'underline': False})
            else:           # underline
                parts.append({'text': inner_text, 'bold': False, 'underline': True})
            last_end = m.end()
        if last_end < len(text):
            parts.append({'text': text[last_end:], 'bold': False, 'underline': False})
        return parts if parts else [{'text': text, 'bold': False, 'underline': False}]

    # ------------------------------------------
    def detect_content_type(self, line):
        line = line.strip()
        if not line:
            return 'empty'
        if '|' in line and len(line.split('|')) > 2:
            return 'table'
        if re.match(r'^#+', line):
            return 'heading'
        if re.search(r'\$\$.*?\$\$|\$.*?\$', line):
            return 'formula'
        if re.match(r'^(شکل|جدول)\s*\d+', line):
            return 'caption'
        return 'text'

    def add_heading(self, line, level=1):
        text = re.sub(r'^#+\s*', '', line)
        text = self.text_processor.clean_text(text)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_rtl(p)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'B Nazanin'
        run.font.size = Pt(18 - level * 2)
        run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')

    def add_formula(self, line):
        formulas = re.findall(r'\$\$.*?\$\$|\$.*?\$', line)
        for f in formulas:
            f = f.strip('$').strip()
            p = self.doc.add_paragraph(f)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.runs[0]
            r.font.name = 'Cambria Math'
            r.font.size = Pt(14)

    def add_caption(self, text):
        text = self.text_processor.clean_text(text)
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_rtl(p)
        for run in p.runs:
            run.bold = True
            run.font.name = 'B Nazanin'
            run.font.size = Pt(13)
            run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')

    def add_table(self, lines):
        rows = []
        for ln in lines:
            if not ln.strip():
                continue
            parts = [self.text_processor.clean_text(p.strip()) for p in ln.strip('|').split('|')]
            if len(parts) > 1:
                rows.append(parts)
        if not rows:
            return
        cols = max(len(r) for r in rows)
        rows = [r + ['']*(cols - len(r)) for r in rows]
        # حذف خط جداکننده markdown
        if len(rows) > 1 and all(set(c.strip()) <= {'-', ':', '|'} for c in rows[1]):
            rows.pop(1)
        if not rows:
            return
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        table.autofit = False
        for i, row_data in enumerate(rows):
            is_header = (i == 0)
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                self._set_cell_borders(cell)
                self._set_cell_shading(cell, is_header)
                self._set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for part in self._parse_inline_marks(cell_data):
                    run = p.add_run(part['text'])
                    if re.search(r'[A-Za-z0-9]', part['text']):
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                    else:
                        run.font.name = 'B Nazanin'
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
                    if part['bold'] or is_header:
                        run.bold = True
                    if part['underline']:
                        run.underline = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._set_rtl(p)
        self.doc.add_paragraph()

    def add_text(self, text):
        text = self.text_processor.clean_text(text)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        self._set_rtl(p)
        for part in self._parse_inline_marks(text):
            run = p.add_run(part['text'])
            run.font.name = 'B Nazanin'
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:cs'), 'B Nazanin')
            run.bold = part['bold']
            run.underline = part['underline']

    def process_text(self, text):
        if not text or not isinstance(text, str):
            self.add_text("⚠️ ورودی خالی یا نامعتبر بود.")
            return
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            ln = lines[i]
            t = self.detect_content_type(ln)
            if t == 'empty':
                i += 1
                continue
            elif t == 'table':
                block = []
                while i < len(lines) and '|' in lines[i]:
                    block.append(lines[i])
                    i += 1
                self.add_table(block)
                continue
            elif t == 'heading':
                level = min(len(re.match(r'^#+', ln).group()), 3)
                self.add_heading(ln, level)
            elif t == 'formula':
                self.add_formula(ln)
            elif t == 'caption':
                self.add_caption(ln)
            else:
                self.add_text(ln)
            i += 1

    def save_to_stream(self):
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf

# ---------------- Flask routes ----------------
@app.route('/generate', methods=['POST'])
def generate_word():
    try:
        data = request.get_json(force=True, silent=True)
        if not data or 'text' not in data:
            return jsonify({'error': 'متن الزامی است'}), 400
        text = data.get('text', '')
        gen = SmartDocumentGenerator()
        gen.process_text(text)
        stream = gen.save_to_stream()
        if not stream.getvalue():
            raise ValueError("خروجی خالی است.")
        return send_file(
            stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='document.docx'
        )
    except Exception as e:
        safe_error = html.escape(str(e).encode('utf-8', 'ignore').decode('utf-8', 'ignore'))
        return jsonify({'error': f'Safe Fail ⛔ {safe_error}'}), 200

@app.route('/')
def home():
    return jsonify({'message': 'Persian DOCX Generator — Stable Mode ✅'})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8001)
